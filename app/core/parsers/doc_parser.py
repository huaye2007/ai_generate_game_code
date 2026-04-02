"""文档解析器 - 解析 Word/Excel 文档转 Markdown，生成需求文档"""
from pathlib import Path
from langchain_core.documents import Document

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class DocParser:
    """解析 Word 文档，转为 Markdown"""

    def parse_files(self, file_paths: list[str]) -> list[Document]:
        """解析多个文档"""
        docs = []
        for fp in file_paths:
            p = Path(fp)
            if p.suffix in (".docx", ".doc"):
                docs.extend(self._parse_docx(p))
            elif p.suffix == ".md":
                docs.extend(self._parse_markdown(p))
            elif p.suffix == ".txt":
                docs.extend(self._parse_text(p))
        return docs

    def _parse_docx(self, file_path: Path) -> list[Document]:
        """解析 Word 文档"""
        if DocxDocument is None:
            return [Document(page_content="python-docx 未安装", metadata={"source": file_path.name})]
        try:
            doc = DocxDocument(str(file_path))
            md_content = self._docx_to_markdown(doc)
            return [Document(
                page_content=md_content,
                metadata={"source": file_path.name, "category": "document", "format": "docx"},
            )]
        except Exception as e:
            return [Document(page_content=f"解析失败: {e}", metadata={"source": file_path.name})]

    def _docx_to_markdown(self, doc) -> str:
        """将 Word 文档转为 Markdown"""
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            elif "list" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # 处理表格
        for table in doc.tables:
            lines.append("")
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        return "\n".join(lines)

    def _parse_markdown(self, file_path: Path) -> list[Document]:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        return [Document(page_content=content, metadata={"source": file_path.name, "category": "document", "format": "md"})]

    def _parse_text(self, file_path: Path) -> list[Document]:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        return [Document(page_content=content, metadata={"source": file_path.name, "category": "document", "format": "txt"})]

    def to_markdown(self, file_path: str) -> str:
        """将文档转为 Markdown 字符串"""
        p = Path(file_path)
        if p.suffix in (".docx", ".doc") and DocxDocument:
            doc = DocxDocument(str(p))
            return self._docx_to_markdown(doc)
        elif p.suffix == ".md":
            return p.read_text(encoding="utf-8", errors="ignore")
        elif p.suffix == ".txt":
            return p.read_text(encoding="utf-8", errors="ignore")
        return f"不支持的文件格式: {p.suffix}"
